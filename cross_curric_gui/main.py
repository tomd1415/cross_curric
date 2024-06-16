import os
import platform
from openai import OpenAI
from xhtml2pdf import pisa
from docx import Document
from config import Config
from event_handler import EventHandler

def convert_html_to_pdf(source_html, output_filename):
    """Convert HTML content to PDF."""
    with open(output_filename, "w+b") as output_file:
        pisa_status = pisa.CreatePDF(source_html, dest=output_file)
    return pisa_status.err

def generate_word_doc(html_content, output_filename):
    """Convert HTML content to Word document."""
    document = Document()

    # Parsing HTML content and adding to document (simplified)
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'li']):
        if element.name == 'h1':
            document.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            document.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            document.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            document.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            document.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                document.add_paragraph(li.get_text(), style='List Bullet')
    
    document.save(output_filename)

def generate_html_report():
    """Generate HTML report using OpenAI API."""
    client = OpenAI(api_key=Config.API_KEY)
    assistant = client.beta.assistants.retrieve(Config.ASSISTANT_ID)
    assistant = client.beta.assistants.update(
        assistant_id=assistant.id,
        tool_resources={"file_search": {"vector_store_ids": [Config.VECTOR_STORE_ID]}}
    )
    thread = client.beta.threads.create()

    with open(Config.HTML_FILENAME, "w") as html_file:
        html_file.write(Config.STYLE)
        event_handler = EventHandler(html_file)
        
        with client.beta.threads.runs.stream(
            thread_id=thread.id,
            assistant_id=assistant.id,
            max_prompt_tokens=2200,
            max_completion_tokens=2500,
            instructions=f"""Please look through the files in the vector store and find cross-curriculum links between {Config.SUBJECT1} and {Config.SUBJECT2}. I would like the report to be organised by year group with a two sentence summary of the links at the start and then details about the links as bullet points below. I would like it grouped by term within year group as in the examples. I would like you to look at the current {Config.SUBJECT2} curriculum as well as the {Config.SUBJECT1} curriculum to find where there could be links in the same term. I would like as many links as possible for each term and I would like it formatted like the examples. Then I would like a short paragraph explaining how more cross-curricular links could be made. This should be about one page per year group in total length. The HTML used to markup the text is not for use in a web pag and therefore only needs the tags that are in the examples. No other tags should be used. Please just show me {Config.YEAR_GROUP} for now. I will come back for the other year groups. Please produce this is a format as close to the examples as you can. Below are the example outputs:""" + Config.PROMPT_EXAMPLES,
            event_handler=event_handler,
        ) as stream:
            stream.until_done()

        html_file.write("</html>")

def ensure_directories():
    """Ensure the output directories exist."""
    base_dir = f"./{Config.SUBJECT1}/{Config.YEAR_GROUP}"
    pdf_dir = os.path.join(base_dir, "pdf")
    docx_dir = os.path.join(base_dir, "docx")
    html_dir = os.path.join(base_dir, "html")

    os.makedirs(pdf_dir, exist_ok=True)
    os.makedirs(docx_dir, exist_ok=True)
    os.makedirs(html_dir, exist_ok=True)

    return pdf_dir, docx_dir, html_dir

def generate_reports_for_all_combinations(subject, year_group):
    """Generate reports for the selected subject combined with all other subjects."""
    subjects = ["Art", "Computing", "Citizenship", "Design and Technology", "English", "French", "Food and Nutrition", "Geography", "History", "Life Skills", "Maths", "Media Studies", "Music", "Physical Education", "Psychology", "PSHE", "Religious Education", "Science"]
    subjects.remove(subject)  # Remove the selected subject from the list
    
    for other_subject in subjects:
        Config.update(subject, other_subject, year_group)

        pdf_dir, docx_dir, html_dir = ensure_directories()

        Config.HTML_FILENAME = os.path.join(html_dir, os.path.basename(Config.HTML_FILENAME))
        Config.PDF_FILENAME = os.path.join(pdf_dir, os.path.basename(Config.PDF_FILENAME))
        Config.DOCX_FILENAME = os.path.join(docx_dir, os.path.basename(Config.DOCX_FILENAME))

        generate_html_report()
        
        with open(Config.HTML_FILENAME, "r") as html_file:
            html_content = html_file.read()

        convert_html_to_pdf(html_content, Config.PDF_FILENAME)
        print(f"Finished! PDF file created: {Config.PDF_FILENAME}")

        generate_word_doc(html_content, Config.DOCX_FILENAME)
        print(f"Finished! Word document created: {Config.DOCX_FILENAME}")
        
        # Optionally open the PDF file
        """
        if platform.system() == 'Darwin':       # macOS
            os.system(f'open "{Config.PDF_FILENAME}"')
        elif platform.system() == 'Windows':    # Windows
            os.startfile(Config.PDF_FILENAME)
        else:                                   # Linux and others
            os.system(f'xdg-open "{Config.PDF_FILENAME}"')
        """

def generate_summary_report(uploaded_files):
    """Generate a summary report from the uploaded HTML files."""
    combined_html_content = ""
    
    for file in uploaded_files:
        with open(file, "r") as f:
            combined_html_content += f.read()
    
    client = OpenAI(api_key=Config.API_KEY)
    assistant = client.beta.assistants.retrieve(Config.ASSISTANT_ID)
    assistant = client.beta.assistants.update(
        assistant_id=assistant.id,
        tool_resources={"file_search": {"vector_store_ids": [Config.VECTOR_STORE_ID]}}
    )
    thread = client.beta.threads.create()

    with open(Config.SUMMARY_HTML_FILENAME, "w") as html_file:
        html_file.write(Config.STYLE)
        event_handler = EventHandler(html_file)
        
        with client.beta.threads.runs.stream(
            thread_id=thread.id,
            assistant_id=assistant.id,
            max_prompt_tokens=2200,
            max_completion_tokens=2500,
            instructions=f"""The audiance for this document is a classroom teacher of the subject. Please summarise the following content and find cross-curriculum links for {Config.SUBJECT1} in {Config.YEAR_GROUP}. The content has come from about 17 PDF files that contain information about cross-curriculum links. I would like the report to be organised by half-term (Autumn 1, Autumn 2, Spring 1, Spring 2, Summer 1, Summer 2), in each half term make it clear which subjects have the cross-curriculm links. The document should start with a three-sentence summary of the cross-curriculum links between {Config.SUBJECT1} and every other subject in this year group. Then details about the links as bullet points below each term heading making it clear which subjects are involved in the cross-curriculm link. Please try and ensure there is a variaty of subjects in the links accross the document.. """ + Config.SUMMARY_PROMPT + combined_html_content,
            event_handler=event_handler,
        ) as stream:
            stream.until_done()

        html_file.write("</html>")

    with open(Config.SUMMARY_HTML_FILENAME, "r") as html_file:
        html_content = html_file.read()

    convert_html_to_pdf(html_content, Config.SUMMARY_PDF_FILENAME)
    generate_word_doc(html_content, Config.SUMMARY_DOCX_FILENAME)
        
def main(single_subject_mode=False, uploaded_files=None):
    """Main function to generate HTML report and convert it to PDF and Word."""
    pdf_dir, docx_dir, html_dir = ensure_directories()

    if uploaded_files:
        generate_summary_report(uploaded_files)
    else:
        Config.HTML_FILENAME = os.path.join(html_dir, os.path.basename(Config.HTML_FILENAME))
        Config.PDF_FILENAME = os.path.join(pdf_dir, os.path.basename(Config.PDF_FILENAME))
        Config.DOCX_FILENAME = os.path.join(docx_dir, os.path.basename(Config.DOCX_FILENAME))


        if single_subject_mode:
            generate_reports_for_all_combinations(Config.SUBJECT1, Config.YEAR_GROUP)
        else:
            generate_html_report()
            
            with open(Config.HTML_FILENAME, "r") as html_file:
                html_content = html_file.read()

            convert_html_to_pdf(html_content, Config.PDF_FILENAME)
            print(f"Finished! PDF file created: {Config.PDF_FILENAME}")

            generate_word_doc(html_content, Config.DOCX_FILENAME)
            print(f"Finished! Word document created: {Config.DOCX_FILENAME}")
            
            # Open the generated PDF file based on the operating system
            if platform.system() == 'Darwin':       # macOS
                os.system(f'open "{Config.PDF_FILENAME}"')
            elif platform.system() == 'Windows':    # Windows
                os.startfile(Config.PDF_FILENAME)
            else:                                   # Linux and others
                os.system(f'xdg-open "{Config.PDF_FILENAME}"')

if __name__ == "__main__":
    main()
